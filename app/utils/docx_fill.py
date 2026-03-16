"""
DOCX template prefilling and PDF generation utilities.
Uses pure Python (reportlab) for PDF generation - no external dependencies like LibreOffice.
Includes signature overlay functionality for final PDF.
"""
import io
import hashlib
from typing import Dict, Any, Optional
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader
from PIL import Image
import boto3

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"

# S3 Configuration for signature download
AWS_REGION = "ap-south-2"
SIGNATURE_BUCKET = "fittbot-uploads"


def sha256_bytes(data: bytes) -> str:
    """Calculate SHA256 hash of bytes data."""
    return hashlib.sha256(data).hexdigest()


def get_template_path(template_name: str = "agreement_new.docx") -> Path:
    template_path = TEMPLATES_DIR / template_name
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")
    return template_path


def _replace_placeholder_in_paragraph(paragraph, placeholder: str, value: str) -> bool:
    full_text = paragraph.text

    if placeholder not in full_text:
        return False

    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            return True

    new_text = full_text.replace(placeholder, value)

    if paragraph.runs:
        first_run = paragraph.runs[0]
        for run in paragraph.runs[1:]:
            run.text = ""
        first_run.text = new_text
        return True

    return False


def _get_ordinal_suffix(day: int) -> str:
    """Get ordinal suffix for a day number (st, nd, rd, th)."""
    if 11 <= day <= 13:
        return "th"
    last_digit = day % 10
    if last_digit == 1:
        return "st"
    elif last_digit == 2:
        return "nd"
    elif last_digit == 3:
        return "rd"
    else:
        return "th"


def _replace_underscore_patterns(paragraph, prefill: Dict[str, Any]) -> None:
    """Replace date-related underscore patterns in paragraph."""
    import re

    day = prefill.get("day", "")
    month = prefill.get("month", "")
    year = prefill.get("year", "2025")

    # Get full paragraph text first
    full_text = paragraph.text

    # ONLY process if this looks like a date line (contains "day" or month names)
    is_date_line = "day" in full_text.lower() or any(m in full_text for m in
        ["January", "February", "March", "April", "May", "June",
         "July", "August", "September", "October", "November", "December"])

    if "_" in full_text and is_date_line:
        new_text = full_text

        # Replace full date pattern "___ day of ___, 2025" with "27th December 2025"
        if day and month:
            day_num = int(day)
            ordinal = _get_ordinal_suffix(day_num)
            # Match pattern: underscores + "day of" + underscores/spaces + ", 2025" or ",2025"
            new_text = re.sub(
                r'_+\s*day\s+of\s+_*\s*,?\s*' + str(year),
                f"{day}{ordinal} {month} {year}",
                new_text,
                flags=re.IGNORECASE
            )
            # Also handle if month name is already there
            new_text = re.sub(
                r'_+\s*day\s+of\s+' + month + r'\s*,?\s*' + str(year),
                f"{day}{ordinal} {month} {year}",
                new_text,
                flags=re.IGNORECASE
            )

        # If text changed, update paragraph
        if new_text != full_text and paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""


def _replace_label_value_paragraphs(doc: Document, prefill: Dict[str, Any]) -> None:
    """Replace underscore placeholders that follow label paragraphs."""
    import re

    if not prefill:
        return

    # Mapping of labels to prefill keys (simplified for new template)
    label_to_prefill = {
        "principal place of business at": "gym_address",
        "principal place of business at:": "gym_address",
    }

    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs) - 1:
        para_text = paragraphs[i].text.strip().lower()

        # Check if this paragraph ends with a label
        for label, prefill_key in label_to_prefill.items():
            if para_text.endswith(label) or para_text.endswith(label + ":"):
                value = prefill.get(prefill_key, "")
                if value:
                    # Look at next 1-2 paragraphs for underscore placeholder
                    for offset in [1, 2]:
                        if i + offset < len(paragraphs):
                            next_para = paragraphs[i + offset]
                            next_text = next_para.text.strip()
                            # Check if it's an underscore placeholder
                            if re.match(r'^[_\s]+$', next_text):
                                # Replace with value
                                if next_para.runs:
                                    next_para.runs[0].text = str(value)
                                    for run in next_para.runs[1:]:
                                        run.text = ""
                                else:
                                    next_para.add_run(str(value))
                                break
                break
        i += 1


def _replace_in_tables(doc: Document, replacements: Dict[str, str], prefill: Dict[str, Any] = None) -> None:
    """Replace placeholders in all tables. Handles label-value table format."""

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            # Do regular replacements in all cells
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        _replace_placeholder_in_paragraph(paragraph, placeholder, value or "")


def _replace_in_document(doc: Document, replacements: Dict[str, str], prefill: Dict[str, Any] = None) -> None:
    """Replace placeholders throughout the document."""
    # First: Replace label-value paragraph pairs (e.g., address after "principal place of business at:")
    if prefill:
        _replace_label_value_paragraphs(doc, prefill)

    for paragraph in doc.paragraphs:
        # Try regex-based underscore pattern replacement for date
        if prefill:
            _replace_underscore_patterns(paragraph, prefill)
            # Replace gym name, address, GST, PAN using regex
            _replace_gym_details_in_paragraph(paragraph, prefill)
        # Then do standard placeholder replacements
        for placeholder, value in replacements.items():
            _replace_placeholder_in_paragraph(paragraph, placeholder, value or "")

    _replace_in_tables(doc, replacements, prefill)

    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    for placeholder, value in replacements.items():
                        _replace_placeholder_in_paragraph(paragraph, placeholder, value or "")

        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    for placeholder, value in replacements.items():
                        _replace_placeholder_in_paragraph(paragraph, placeholder, value or "")


def build_placeholder_map(prefill: Dict[str, Any]) -> Dict[str, str]:
    """Build placeholder replacements for date and AND section fields."""
    replacements = {}

    # Date replacements handled by _replace_underscore_patterns now
    # Format: "27th December 2025" instead of "27 day of December, 2025"

    return replacements


def _replace_gym_details_in_paragraph(paragraph, prefill: Dict[str, Any]) -> None:
    """Replace gym name, address, GST, and PAN in paragraph using regex."""
    import re

    if not prefill:
        return

    gym_name = prefill.get("gym_name", "")
    gym_address = prefill.get("gym_address", "")
    gst_number = prefill.get("gst_number", "")
    pan_number = prefill.get("pan_number", "")

    full_text = paragraph.text
    new_text = full_text

    # Replace gym name and add address: pattern like "________, a fitness center having its principal place of business at:"
    if gym_name and "a fitness center" in full_text.lower():
        # Replace underscores before "a fitness center" with gym name
        new_text = re.sub(r'_+\s*,?\s*a fitness center', f'{gym_name}, a fitness center', new_text, flags=re.IGNORECASE)

        # If this line also contains "principal place of business at:" - add address after it
        if gym_address and "principal place of business at" in new_text.lower():
            # Check if line ends with "at:" (possibly with trailing spaces/underscores)
            if re.search(r'at:\s*_*\s*$', new_text, flags=re.IGNORECASE):
                # Replace ending "at:" or "at:___" with "at:\n<address>"
                new_text = re.sub(r'at:\s*_*\s*$', f'at:\n{gym_address}', new_text, flags=re.IGNORECASE)
            elif new_text.strip().endswith("at:"):
                new_text = new_text.rstrip() + f"\n{gym_address}"

    # Replace address: if this line is just underscores (address placeholder line)
    if gym_address:
        if re.match(r'^[_\s]+$', full_text.strip()):
            new_text = gym_address

    # Replace GSTN: handle "GSTN: ___" pattern
    if gst_number and "GSTN:" in full_text.upper():
        new_text = re.sub(r'GSTN:\s*_+', f'GSTN: {gst_number}', new_text, flags=re.IGNORECASE)

    # Replace PAN: handle "PAN: ___" pattern including trailing underscores
    if pan_number and "PAN:" in full_text.upper():
        # Replace PAN: followed by underscores
        new_text = re.sub(r'PAN:\s*_+', f'PAN: {pan_number}', new_text, flags=re.IGNORECASE)
        # Also remove any trailing underscores after the PAN number (before parenthesis or space)
        new_text = re.sub(rf'(PAN:\s*{re.escape(pan_number)})_+', r'\1', new_text, flags=re.IGNORECASE)

    # Update paragraph if changed
    if new_text != full_text and paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def generate_prefilled_docx(prefill: Dict[str, Any], template_name: str = "agreement_new.docx") -> bytes:
    template_path = get_template_path(template_name)
    doc = Document(str(template_path))
    replacements = build_placeholder_map(prefill)
    _replace_in_document(doc, replacements, prefill)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _extract_text_from_docx(docx_bytes: bytes) -> list:
    """Extract paragraphs and tables from DOCX for PDF generation."""
    doc = Document(io.BytesIO(docx_bytes))
    content = []

    for element in doc.element.body:
        if element.tag.endswith('p'):
            # Paragraph
            for para in doc.paragraphs:
                if para._element == element:
                    text = para.text.strip()
                    if text:
                        # Check if it's a heading style
                        is_heading = para.style and 'Heading' in para.style.name if para.style else False
                        is_bold = any(run.bold for run in para.runs if run.bold)
                        content.append({
                            'type': 'paragraph',
                            'text': text,
                            'is_heading': is_heading or is_bold
                        })
                    break
        elif element.tag.endswith('tbl'):
            # Table
            for table in doc.tables:
                if table._element == element:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = '\n'.join(p.text for p in cell.paragraphs).strip()
                            row_data.append(cell_text)
                        table_data.append(row_data)
                    if table_data:
                        content.append({
                            'type': 'table',
                            'data': table_data
                        })
                    break

    return content


def _create_witness_section(prefill: Dict[str, Any], styles) -> list:
    """Create the signature section with two-column layout like DS Zerodha format.
    Left side: Gym name (big), Owner name, Signature placeholder
    Right side: NFCTECH Fitness Private Limited (big) with digital signature details
    """
    from reportlab.platypus import Paragraph, Table, TableStyle, Spacer
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from datetime import datetime

    gym_name = prefill.get("gym_name", "") if prefill else "________________"
    authorized_person = prefill.get("authorized_person", "") if prefill else "________________"
    nfctech_company = "NFCTech Fitness Private Limited"

    # Get current date for digital signature
    now = datetime.now()
    sign_date = now.strftime("%Y-%m-%d %H:%M:%S+05:30")

    # Styles for witness section
    normal_style = ParagraphStyle('WitnessNormal', parent=styles['Normal'], fontSize=10, spaceAfter=3, leading=14)

    elements = []

    elements.append(Spacer(1, 40))

    # Two-column signatory section
    # Left column: Gym name (big), Owner name, Signature space
    left_content = f"""<b><font size="14">{gym_name}</font></b><br/><br/>
<b>{authorized_person}</b><br/><br/>
<br/><br/><br/><br/>
_________________________<br/>
<font size="8">(Authorized Signatory)</font>"""

    # Right column: NFCTech with digital signature format (like DS Zerodha)
    right_content = f"""<b><font size="14">{nfctech_company}</font></b><br/><br/>
<font size="8">Digitally signed by {nfctech_company}<br/>
Reason: Signed by Auth. signatore {nfctech_company}<br/>
Location: Bangalore<br/>
Date: {sign_date}</font>"""

    left_para = Paragraph(left_content, normal_style)
    right_para = Paragraph(right_content, normal_style)

    # Create two-column table
    col_width = (A4[0] - 1.5*inch) / 2
    witness_table = Table([[left_para, right_para]], colWidths=[col_width, col_width])
    witness_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'LEFT'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))

    elements.append(witness_table)

    return elements


def convert_docx_to_pdf(docx_bytes: bytes, prefill: Dict[str, Any] = None) -> bytes:
    """Convert DOCX bytes to PDF using pure Python (reportlab)."""
    content = _extract_text_from_docx(docx_bytes)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )

    styles = getSampleStyleSheet()

    # Custom styles
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=6
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        leading=16,
        spaceAfter=10,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=14,
        leading=18,
        spaceAfter=20,
        alignment=1,  # Center
        fontName='Helvetica-Bold'
    )

    # Get values for bold formatting
    gym_name = prefill.get("gym_name", "") if prefill else ""
    gym_address = prefill.get("gym_address", "") if prefill else ""
    gst_number = prefill.get("gst_number", "") if prefill else ""
    pan_number = prefill.get("pan_number", "") if prefill else ""
    day = prefill.get("day", "") if prefill else ""
    month = prefill.get("month", "") if prefill else ""

    story = []
    skip_signature_content = False

    for item in content:
        if item['type'] == 'paragraph':
            text = item['text']

            # Detect signature section at end - add custom formatted section
            # Look for patterns that indicate the signature block
            if any(marker in text.upper() for marker in ['IN WITNESS WHEREOF', 'FOR THE GYM PARTNER', 'SIGNATURE']):
                if not skip_signature_content:
                    skip_signature_content = True
                    # Add custom two-column signature section
                    witness_elements = _create_witness_section(prefill, styles)
                    story.extend(witness_elements)
                continue

            # Skip content after signature section starts
            if skip_signature_content:
                continue

            # Escape special XML characters
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

            # Make full date bold (e.g., "27th December 2025")
            if day and month:
                year = prefill.get("year", "2025") if prefill else "2025"
                day_num = int(day)
                ordinal = _get_ordinal_suffix(day_num)
                full_date = f"{day}{ordinal} {month} {year}"
                if full_date in text:
                    text = text.replace(full_date, f"<b>{full_date}</b>")

            # Make gym name bold
            if gym_name and gym_name in text:
                text = text.replace(gym_name, f"<b>{gym_name}</b>")

            # Make gym address bold
            if gym_address and gym_address in text:
                text = text.replace(gym_address, f"<b>{gym_address}</b>")

            # Make GST number bold
            if gst_number and gst_number in text:
                text = text.replace(gst_number, f"<b>{gst_number}</b>")

            # Make PAN number bold
            if pan_number and pan_number in text:
                text = text.replace(pan_number, f"<b>{pan_number}</b>")

            if item.get('is_heading'):
                # Check if it looks like a title
                if len(text) < 50 and text.isupper():
                    story.append(Paragraph(text, title_style))
                else:
                    story.append(Paragraph(text, heading_style))
            else:
                story.append(Paragraph(text, normal_style))

        elif item['type'] == 'table':
            if skip_signature_content:
                continue

            table_data = item['data']
            if table_data:
                # Calculate column widths
                num_cols = max(len(row) for row in table_data)
                available_width = A4[0] - 1.5*inch
                col_width = available_width / num_cols

                # Process table data
                table_style = ParagraphStyle('TableCell', parent=styles['Normal'], fontSize=9)
                processed_data = []
                for row in table_data:
                    processed_row = []
                    for cell in row:
                        cell_text = cell.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        processed_row.append(Paragraph(cell_text, table_style))
                    processed_data.append(processed_row)

                # Create table with Paragraph objects
                t = Table(processed_data, colWidths=[col_width]*num_cols)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(Spacer(1, 10))
                story.append(t)
                story.append(Spacer(1, 10))

    # Always add signature section at the end if not already added
    if not skip_signature_content:
        witness_elements = _create_witness_section(prefill, styles)
        story.extend(witness_elements)

    if story:
        doc.build(story)
    else:
        # If no content extracted, create a simple placeholder PDF
        c = canvas.Canvas(buffer, pagesize=A4)
        c.drawString(100, 750, "Agreement Document")
        c.save()

    return buffer.getvalue()


def generate_prefilled_pdf(prefill: Dict[str, Any], template_name: str = "agreement_new.docx") -> bytes:
    """Generate PDF from prefilled DOCX template using pure Python."""
    docx_bytes = generate_prefilled_docx(prefill, template_name)
    pdf_bytes = convert_docx_to_pdf(docx_bytes, prefill)
    return pdf_bytes


def download_signature_from_s3(signature_url: str) -> bytes:

    parsed = urlparse(signature_url.split('?')[0])  # Remove query params like ?v=timestamp

    # Extract key from path (remove leading /)
    s3_key = parsed.path.lstrip('/')

    # Download from S3
    s3 = boto3.client("s3", region_name=AWS_REGION)
    response = s3.get_object(Bucket=SIGNATURE_BUCKET, Key=s3_key)
    return response['Body'].read()


def add_signature_to_pdf(
    pdf_bytes: bytes,
    signature_bytes: bytes,
    gym_name: str,
    authorized_person: Optional[str] = None,
    nfctech_name: str = "Nishad Shaik",
    nfctech_designation: str = "Director"
) -> bytes:
    """
    Add gym owner signature to the existing IN WITNESS section.
    Signature is placed on the LEFT side (gym owner's column).
    Only overlays the signature image - no duplicate text.
    """

    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()

    # Copy all pages except the last one
    for i in range(len(reader.pages) - 1):
        writer.add_page(reader.pages[i])

    # Get the last page
    last_page = reader.pages[-1]
    page_width = float(last_page.mediabox.width)
    page_height = float(last_page.mediabox.height)

    # Load and resize signature image
    sig_image = Image.open(io.BytesIO(signature_bytes))

    # Convert to RGBA if needed
    if sig_image.mode != 'RGBA':
        sig_image = sig_image.convert('RGBA')

    # Resize signature - smaller size to fit above the line
    max_sig_width = 120
    sig_width, sig_height = sig_image.size
    if sig_width > max_sig_width:
        ratio = max_sig_width / sig_width
        sig_width = max_sig_width
        sig_height = int(sig_height * ratio)
        sig_image = sig_image.resize((sig_width, sig_height), Image.LANCZOS)

    # Limit max height
    max_sig_height = 60
    if sig_height > max_sig_height:
        ratio = max_sig_height / sig_height
        sig_height = max_sig_height
        sig_width = int(sig_width * ratio)
        sig_image = sig_image.resize((sig_width, sig_height), Image.LANCZOS)

    # Create overlay PDF with just the signature
    packet = io.BytesIO()
    c = canvas.Canvas(packet, pagesize=(page_width, page_height))

    # LEFT column for gym owner signature (about 60px from left margin)
    left_column_x = 60
    sig_y = 480 # Position above the signature line at bottom

    # Draw the signature image
    sig_buffer = io.BytesIO()
    sig_image.save(sig_buffer, format='PNG')
    sig_buffer.seek(0)

    c.drawImage(
        ImageReader(sig_buffer),
        left_column_x,
        sig_y,
        width=sig_width,
        height=sig_height,
        mask='auto'  # Handle transparency
    )

    c.save()
    packet.seek(0)

    # Merge overlay onto last page
    overlay = PdfReader(packet)
    if overlay.pages:
        last_page.merge_page(overlay.pages[0])

    writer.add_page(last_page)

    # Write final PDF
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def generate_prefilled_pdf_with_signature(
    prefill: Dict[str, Any],
    signature_url: str,
    template_name: str = "agreement_new.docx"
) -> bytes:
    """
    Generate PDF with signature section containing both signatories:
    - Gym Owner (left side): Name from prefill, signature from S3
    - NFCTech Fitness Private Limited (right side): Company name and Nishad Shaik
    """

    # Generate base PDF
    pdf_bytes = generate_prefilled_pdf(prefill, template_name)

    # Download signature
    signature_bytes = download_signature_from_s3(signature_url)

    # Get signatory info from prefill
    gym_name = prefill.get("gym_name", "")
    authorized_person = prefill.get("authorized_person")
    nfctech_name = prefill.get("nfctech_name", "Nishad Shaik")
    nfctech_designation = prefill.get("nfctech_designation", "Director")

    final_pdf = add_signature_to_pdf(
        pdf_bytes,
        signature_bytes,
        gym_name,
        authorized_person,
        nfctech_name,
        nfctech_designation
    )

    return final_pdf
